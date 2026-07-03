# -*- coding: utf-8 -*-
"""Экспорт DOCX-отчёта: титул, сводка, диагнозы, топ-5 гипотез таблицей,
«Почему не предложено», дорожная карта отдельным разделом, карточки гипотез."""
from __future__ import annotations

import io
from datetime import date

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..diagnostics import DiagnosticsResult
from ..models import Hypothesis, Project, RoadmapItem, TailingsReport

STAGE_LABELS = {"lab": "Лаборатория", "pilot": "ОПИ", "rollout": "Тираж"}
CAPEX_LABELS = {"low": "низкий", "med": "средний", "medium": "средний", "high": "высокий"}


def build_report_docx(project: Project, report: TailingsReport, diag: DiagnosticsResult,
                      hypotheses: list[Hypothesis],
                      roadmap: list[RoadmapItem] | None = None) -> bytes:
    d = docx.Document()

    # титул
    title = d.add_heading("Фабрика гипотез — отчёт", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = d.add_paragraph(f"Проект: {project.name or project.plant}")
    p.add_run(f"\nЦель: {project.goal or '—'}")
    p.add_run(f"\nТип хвостов: {report.tail_type}")
    p.add_run(f"\nДата: {date.today().isoformat()}")

    # сводка
    d.add_heading("Сводка по отчёту института", level=1)
    ni_t = report.losses_tonnes.get("Ni")
    cu_t = report.losses_tonnes.get("Cu")
    rec_ni = report.recoverable_total.get("Ni")
    d.add_paragraph(
        f"Отвальные хвосты: {report.tails_tonnes:,.0f} СМТ. "
        f"Потери: Ni {ni_t:,.1f} т ({report.grade.get('Ni', 0):.4f}%), "
        f"Cu {cu_t:,.1f} т ({report.grade.get('Cu', 0):.4f}%). "
        f"Извлекаемый потенциал: Ni {rec_ni:,.1f} т "
        f"({report.recoverable_pct.get('Ni', 0):.1f}%), "
        f"Cu {report.recoverable_total.get('Cu', 0):,.1f} т "
        f"({report.recoverable_pct.get('Cu', 0):.1f}%).".replace(",", " "))

    d.add_heading("Диагнозы", level=1)
    for dg in diag.diagnoses:
        d.add_paragraph(f"[{dg.rule_id}/{dg.element}] {dg.title}", style="List Bullet")
        d.add_paragraph(dg.text)

    if diag.not_proposed:
        d.add_heading("Почему не предложено (неизвлекаемые формы)", level=1)
        for x in diag.not_proposed:
            d.add_paragraph(
                f"{x['element']} · {x['form']}: {x['tonnes']:,.1f} т — {x['reason']}"
                .replace(",", " "), style="List Bullet")

    # топ-5 таблицей
    d.add_heading("Топ-5 гипотез", level=1)
    top = hypotheses[:5]
    table = d.add_table(rows=1 + len(top), cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, name in enumerate(["№", "Гипотеза", "Передел", "Эффект, т/год",
                              "Эффект, $/год", "CAPEX"]):
        hdr[i].text = name
    for n, h in enumerate(top, 1):
        row = table.rows[n].cells
        row[0].text = str(n)
        row[1].text = h.title
        row[2].text = h.process_area
        row[3].text = f"{h.effect.tonnes_expected:,.0f}".replace(",", " ")
        row[4].text = f"{h.effect.money_usd:,.0f}".replace(",", " ")
        row[5].text = CAPEX_LABELS.get(str(h.feasibility.get("capex", "med")).lower(), "средний")

    # дорожная карта
    if roadmap:
        d.add_heading("Дорожная карта проверки", level=1)
        table = d.add_table(rows=1 + len(roadmap), cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, name in enumerate(["Гипотеза", "Стадия", "Начало", "Конец",
                                  "Ресурс", "Критерий ворот"]):
            hdr[i].text = name
        for n, it in enumerate(roadmap, 1):
            row = table.rows[n].cells
            row[0].text = it.hypothesis_title[:60]
            row[1].text = STAGE_LABELS.get(it.stage, it.stage)
            row[2].text = it.start
            row[3].text = it.end
            row[4].text = it.resource or ""
            row[5].text = (it.gate_criterion or "")[:80]

    # карточки гипотез
    d.add_heading("Карточки гипотез", level=1)
    for n, h in enumerate(hypotheses, 1):
        d.add_heading(f"{n}. {h.title}", level=2)
        meta = d.add_paragraph()
        meta.add_run(f"Передел: {h.process_area} · Элемент: {h.element} · "
                     f"Статус: {h.status} · Score: {h.score:.3f}").font.size = Pt(9)
        d.add_paragraph(f"Механизм: {h.mechanism}")
        d.add_paragraph(
            f"Эффект: до {h.effect.tonnes_max:,.0f} т/год, ожидаемо "
            f"{h.effect.tonnes_expected:,.0f} т/год ≈ ${h.effect.money_usd:,.0f}. "
            f"Допущения: {h.effect.assumptions}".replace(",", " "))
        if h.rationale:
            d.add_paragraph("Обоснование из литературы:")
            for c in h.rationale:
                mark = "✓" if c.verified else "⚠ требует проверки"
                d.add_paragraph(f"«{c.quote}» — {c.source or c.chunk_id}"
                                + (f", с. {c.page}" if c.page else "") + f" [{mark}]",
                                style="List Bullet")
        if h.risks:
            d.add_paragraph("Риски: " + "; ".join(h.risks))
        if h.verification_plan:
            d.add_paragraph("План проверки:")
            for s in h.verification_plan:
                d.add_paragraph(
                    f"{s.n}. {s.action} ({s.duration}). Успех: {s.success_criterion}. "
                    f"Провал: {s.fail_criterion}", style="List Number")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
