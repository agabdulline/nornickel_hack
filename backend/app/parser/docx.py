# -*- coding: utf-8 -*-
"""Парсер docx с эталонными гипотезами экспертов (ground truth для eval).

Формат файлов кейса: заголовок-абзац + таблица в одну колонку,
каждая строка — «N. Текст гипотезы». Парсим и таблицы, и абзацы.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

import docx

_ITEM_RE = re.compile(r"^\s*(\d+)\s*[.)]\s*(.+)", re.DOTALL)


def parse_expert_hypotheses(source, source_name: str = "") -> list[dict]:
    """-> [{"n": 1, "title": "...", "text": "..."}] в порядке нумерации."""
    if isinstance(source, (str, Path)):
        document = docx.Document(str(source))
    elif isinstance(source, bytes):
        document = docx.Document(io.BytesIO(source))
    else:
        document = docx.Document(source)

    texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            seen_ids = set()
            for cell in row.cells:  # merged-ячейки повторяются — дедуп по id
                if id(cell._tc) in seen_ids:
                    continue
                seen_ids.add(id(cell._tc))
                if cell.text.strip():
                    texts.append(cell.text.strip())
    for p in document.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())

    items: list[dict] = []
    seen_n = set()
    for t in texts:
        m = _ITEM_RE.match(t)
        if not m:
            continue
        n = int(m.group(1))
        if n in seen_n or n > 50:
            continue
        seen_n.add(n)
        body = m.group(2).strip()
        lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
        items.append({
            "n": n,
            "title": lines[0] if lines else body,
            "text": body,
        })
    items.sort(key=lambda x: x["n"])
    return items


# --- эталонные гипотезы по фабрикам (для novelty-бейджа и few-shot) ---

_PLANT_PATTERNS = {
    "КГМК": "кгмк",
    "НОФ вкр": "ноф вкр",
    "НОФ мед": "ноф мед",
    "ТОФ": "тоф",
}


def _find_expert_docx(pattern: str):
    import os
    import unicodedata
    from ..config import DATA_CASE
    if not DATA_CASE.exists():
        return None
    for root, _dirs, files in os.walk(DATA_CASE):
        for f in files:
            nf = unicodedata.normalize("NFC", f).lower()
            if nf.startswith("гипотезы") and nf.endswith(".docx") \
                    and pattern in nf.replace("_", " "):
                return f"{root}/{f}"
    return None


def expert_titles_for_plant(plant: str) -> list[str]:
    """Эталонные гипотезы экспертов ЭТОЙ фабрики (novelty / бейдж совпадения)."""
    low = (plant or "").lower()
    key = ("КГМК" if "кгмк" in low else
           "НОФ вкр" if "вкрапл" in low else
           "НОФ мед" if "мед" in low else
           "ТОФ" if "тоф" in low else None)
    if not key:
        return []
    path = _find_expert_docx(_PLANT_PATTERNS[key])
    if not path:
        return []
    try:
        return [x["title"] for x in parse_expert_hypotheses(path)]
    except Exception:  # noqa: BLE001 — отсутствие эталонов не должно ронять пайплайн
        return []


def cross_plant_examples(plant: str, per_plant: int = 2) -> list[str]:
    """Гипотезы экспертов ДРУГИХ фабрик — few-shot образец конкретности для
    генератора. Эталоны текущей фабрики исключаются (никакого лика ground truth)."""
    low = (plant or "").lower()
    current = ("КГМК" if "кгмк" in low else
               "НОФ" if ("ноф" in low or "вкрапл" in low or "мед" in low or "норильск" in low)
               else "ТОФ" if ("тоф" in low or "талнах" in low) else None)
    out: list[str] = []
    for key, pattern in _PLANT_PATTERNS.items():
        if current and key.startswith(current):
            continue
        path = _find_expert_docx(pattern)
        if not path:
            continue
        try:
            items = parse_expert_hypotheses(path)
        except Exception:  # noqa: BLE001
            continue
        out.extend(x["title"] for x in items[:per_plant])
    return out[:6]
