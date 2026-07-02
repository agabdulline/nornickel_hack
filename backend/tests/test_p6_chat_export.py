# -*- coding: utf-8 -*-
"""P6: чат-интерпретатор (контекст-сборщик + мок LLM) и экспорт DOCX/CSV/JSON."""
import io
import json

import pytest

from backend.app.chat import answer, build_context
from backend.app.diagnostics import run_diagnostics
from backend.app.kb.index import KBIndex
from backend.app.kb.textnorm import chunk_pages
from backend.app.models import Citation, Effect, Hypothesis, Project, Step
from backend.app.parser.recover import recover
from backend.app.parser.xlsx import parse_workbook
from backend.tests.conftest import find_case_file, requires_data


@pytest.fixture(scope="module")
def ctx_example2():
    res = parse_workbook(find_case_file(r"Пример 2/Хвосты.*Вкр\.xlsx$"))
    r = res.reports[0]
    recover(r, llm=None)
    diag = run_diagnostics(r)
    hyps = [Hypothesis(id="h01-aaaaaa", title="Замена насадок гидроциклонов", score=0.8,
                       process_area="классификация",
                       effect=Effect(tonnes_max=845.7, tonnes_expected=211.4, money_usd=3_488_000),
                       rationale=[Citation(quote="ц", chunk_id="x", verified=True)]),
            Hypothesis(id="h02-bbbbbb", title="Футеровка мельниц", score=0.6,
                       process_area="измельчение",
                       effect=Effect(tonnes_max=1106, tonnes_expected=243, money_usd=4_016_000))]
    project = Project(id="p1", plant="НОФ (вкрапленные)")
    return r, diag, hyps, project


@requires_data
def test_context_builder_includes_cell_and_rule(ctx_example2, tmp_path):
    """Обязательный тест 8.1: для «откуда 846 т» контекст содержит ячейку
    +125/закрытый Pnt/Cp и правило R1."""
    report, diag, hyps, project = ctx_example2
    kb = KBIndex(root=tmp_path, use_dense=False)
    ctx = build_context("откуда 846 т закрытого никеля?", report, diag, hyps, project, kb)

    top_keys = [c["ячейка"] for c in ctx["отчёт"]["топ_ячейки_потерь"]]
    assert "+125/Закрытый Pnt/Cp/Ni" in top_keys
    r1 = [d for d in ctx["диагнозы"] if d["правило"] == "R1"]
    assert r1 and "+125/Закрытый Pnt/Cp/Ni" in r1[0]["ячейки"]
    assert ctx["веса_ранжирования"]["money"] == 0.4
    assert any(h["id"] == "h01-aaaaaa" for h in ctx["гипотезы_кратко"])


@requires_data
def test_chat_answer_with_mock_llm(ctx_example2, tmp_path):
    report, diag, hyps, project = ctx_example2
    kb = KBIndex(root=tmp_path, use_dense=False)

    class FakeLLM:
        enabled = True
        def chat(self, messages, **kw):
            # проверяем, что контекст дошёл до модели
            joined = json.dumps([m["content"] for m in messages], ensure_ascii=False)
            assert "845.7" in joined or "845.73" in joined
            return {"content": json.dumps({
                "text": "845.73 т — это закрытый Pnt/Cp в классе +125 [+125/Закрытый Pnt/Cp/Ni], "
                        "диагноз недоизмельчения [R1].",
                "references": [{"type": "cell", "id": "+125/Закрытый Pnt/Cp/Ni"},
                               {"type": "rule", "id": "R1"}]}, ensure_ascii=False), "usage": {}}

    ans = answer("откуда 845.73 т?", [], report, diag, hyps, project, kb, llm=FakeLLM())
    assert "R1" in ans.text
    types = {r.type for r in ans.references}
    assert types == {"cell", "rule"}


@requires_data
def test_chat_offline_fallback(ctx_example2, tmp_path):
    report, diag, hyps, project = ctx_example2
    kb = KBIndex(root=tmp_path, use_dense=False)

    class DeadLLM:
        enabled = False
        def chat(self, *a, **kw):
            from backend.app.llm import LLMUnavailable
            raise LLMUnavailable("нет ключа")

    ans = answer("почему?", [], report, diag, hyps, project, kb, llm=DeadLLM())
    assert ans.references and ans.references[0].type == "rule"
    assert "LLM недоступна" in ans.text


@requires_data
def test_docx_export(ctx_example2):
    import docx as docxlib
    from datetime import date
    from backend.app.export.report_docx import build_report_docx
    from backend.app.hypotheses.roadmap import build_roadmap

    report, diag, hyps, project = ctx_example2
    for h in hyps:
        h.status = "accepted"
        h.hypothesis_type = "classification"
        h.verification_plan = [Step(n=1, action="ОПИ", duration="4 нед",
                                    success_criterion="-20%", fail_criterion="рост")]
    roadmap = build_roadmap(hyps, start=date(2026, 7, 6))
    data = build_report_docx(project, report, diag, hyps, roadmap)
    assert data[:2] == b"PK"
    d = docxlib.Document(io.BytesIO(data))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Фабрика гипотез" in text
    assert "Почему не предложено" in text
    assert len(d.tables) >= 2  # топ-5 + дорожная карта


@requires_data
def test_csv_and_json_export(ctx_example2):
    from datetime import date
    from backend.app.export.tasks import to_project_json, to_tasks_csv
    from backend.app.hypotheses.roadmap import build_roadmap

    report, diag, hyps, project = ctx_example2
    for h in hyps:
        h.status = "accepted"
        h.hypothesis_type = "liner"
    roadmap = build_roadmap(hyps, start=date(2026, 7, 6))
    csv_text = to_tasks_csv(hyps, roadmap)
    header = csv_text.splitlines()[0]
    for col in ("stage", "start", "end", "resource", "criterion"):
        assert col in header
    assert "Футеровка мельниц" in csv_text

    payload = to_project_json(project, report, diag, hyps, roadmap)
    dumped = json.dumps(payload, ensure_ascii=False)
    assert payload["project"]["plant"] == "НОФ (вкрапленные)"
    assert len(payload["roadmap"]) == len(roadmap)
    # ключ LLM не должен утекать ни в один экспорт
    from backend.app.config import settings
    if settings.has_key:
        assert settings.llm_api_key not in dumped
        assert settings.llm_api_key not in csv_text
